import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from typing import List, Dict, Any
from pathlib import Path
from core.logger import setup_logger

logger = setup_logger("exporter")

class DataExporter:
    """Класс для экспорта данных"""
    
    def export_to_excel(self, data: List[Dict[str, Any]], filename: str):
        """Экспорт в Excel"""
        
        try:
            df = pd.DataFrame(data)
            df.to_excel(filename, index=False, engine='openpyxl')
            logger.info(f"Экспорт в Excel: {filename}")
        except Exception as e:
            logger.error(f"Ошибка экспорта в Excel: {e}")
            raise
    
    def export_to_word(self, data: List[Dict[str, Any]], filename: str):
        """Экспорт в Word"""
        
        try:
            doc = Document()
            doc.add_heading('Результаты парсинга', 0)
            
            for i, item in enumerate(data, 1):
                doc.add_heading(f'Запись {i}', level=1)
                
                for key, value in item.items():
                    if value:
                        doc.add_paragraph(f'{key}: {value}', style='List Bullet')
                
                doc.add_page_break()
            
            doc.save(filename)
            logger.info(f"Экспорт в Word: {filename}")
            
        except Exception as e:
            logger.error(f"Ошибка экспорта в Word: {e}")
            raise
    
    def export_to_pdf(self, data: List[Dict[str, Any]], filename: str):
        """Экспорт в PDF"""
        
        try:
            c = canvas.Canvas(filename, pagesize=letter)
            width, height = letter
            
            c.setFont("Helvetica", 16)
            c.drawString(50, height - 50, "Результаты парсинга")
            
            y = height - 100
            c.setFont("Helvetica", 10)
            
            for i, item in enumerate(data, 1):
                if y < 50:
                    c.showPage()
                    y = height - 50
                    c.setFont("Helvetica", 10)
                
                c.drawString(50, y, f"Запись {i}")
                y -= 15
                
                for key, value in item.items():
                    if y < 50:
                        c.showPage()
                        y = height - 50
                        c.setFont("Helvetica", 10)
                    
                    if value:
                        text = f"{key}: {str(value)[:100]}"
                        c.drawString(70, y, text)
                        y -= 12
                
                y -= 20
            
            c.save()
            logger.info(f"Экспорт в PDF: {filename}")
            
        except Exception as e:
            logger.error(f"Ошибка экспорта в PDF: {e}")
            raise